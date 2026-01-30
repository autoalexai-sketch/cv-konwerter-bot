"""
Генератор шаблонов CV для Premium пользователей
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path
from datetime import datetime
from typing import Dict, List


class CVGenerator:
    """Генератор резюме в формате DOCX"""
    
    def __init__(self):
        self.output_dir = Path(__file__).parent / 'generated'
        self.output_dir.mkdir(exist_ok=True)
    
    def _set_font(self, run, font_name='Calibri', font_size=None, bold=False, italic=False, color=None):
        """Установить шрифт с поддержкой Unicode (кириллица, польские диакритики)"""
        run.font.name = font_name
        if font_size:
            run.font.size = font_size
        if bold:
            run.font.bold = True
        if italic:
            run.font.italic = True
        if color:
            run.font.color.rgb = color
    
    def _apply_calibri_to_all(self, doc: Document):
        """Применить шрифт Calibri ко всему документу для корректного отображения кириллицы"""
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                if not run.font.name:  # Только если шрифт не установлен явно
                    run.font.name = 'Calibri'
        
        # Применяем к таблицам (если есть)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            if not run.font.name:
                                run.font.name = 'Calibri'
    
    def _add_heading(self, doc: Document, text: str, level: int = 1):
        """Добавить заголовок с форматированием"""
        heading = doc.add_heading(text, level=level)
        heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        return heading
    
    def _add_contact_info(self, doc: Document, data: Dict):
        """Добавить контактную информацию"""
        # Имя и фамилия
        name = doc.add_paragraph()
        name_run = name.add_run(f"{data['imie']} {data['nazwisko']}")
        self._set_font(name_run, font_size=Pt(24), bold=True, color=RGBColor(31, 78, 121))
        name.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Контактные данные
        contact = doc.add_paragraph()
        # Формируем адрес
        address_parts = []
        if 'adres' in data and data['adres']:
            address_parts.append(data['adres'])
        if 'kod_pocztowy' in data and data['kod_pocztowy']:
            address_parts.append(data['kod_pocztowy'])
        if 'miasto' in data and data['miasto']:
            address_parts.append(data['miasto'])
        
        full_address = ', '.join(address_parts) if address_parts else data.get('miasto', '')
        
        contact_text = f"📧 {data['email']}  |  📱 {data['telefon']}  |  📍 {full_address}"
        contact_run = contact.add_run(contact_text)
        self._set_font(contact_run, font_size=Pt(11))
        contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()  # Пустая строка
    
    def _add_section(self, doc: Document, title: str, icon: str = ""):
        """Добавить секцию с иконкой"""
        section = doc.add_paragraph()
        if icon:
            icon_run = section.add_run(f"{icon} ")
            self._set_font(icon_run)
        section_title = section.add_run(title.upper())
        self._set_font(section_title, font_size=Pt(14), bold=True, color=RGBColor(31, 78, 121))
        
        # Линия под заголовком
        doc.add_paragraph('_' * 80)
    
    def generate_klasyczny(self, data: Dict) -> Path:
        """Сгенерировать традиционное CV"""
        doc = Document()
        
        # Настройка отступов
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)
        
        # Контактная информация
        self._add_contact_info(doc, data)
        
        # Stanowisko
        if 'stanowisko' in data:
            stanowisko = doc.add_paragraph()
            stanowisko_run = stanowisko.add_run(data['stanowisko'])
            self._set_font(stanowisko_run, font_size=Pt(16), bold=True)
            stanowisko.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph()
        
        # O sobie
        if 'o_sobie' in data and data['o_sobie']:
            self._add_section(doc, 'O mnie', '👤')
            p_about = doc.add_paragraph()
            about_run = p_about.add_run(data['o_sobie'])
            self._set_font(about_run)
            doc.add_paragraph()
        
        # Doświadczenie zawodowe
        if 'doswiadczenie' in data and data['doswiadczenie']:
            self._add_section(doc, 'Doświadczenie zawodowe', '💼')
            for exp in data['doswiadczenie']:
                p = doc.add_paragraph()
                # Stanowisko i firma
                job_title = p.add_run(f"{exp['stanowisko']} - {exp['firma']}\n")
                self._set_font(job_title, font_size=Pt(12), bold=True)
                # Okres
                period = p.add_run(f"{exp['okres']}\n")
                self._set_font(period, font_size=Pt(10), italic=True)
                # Opis
                if 'opis' in exp and exp['opis']:
                    opis_run = p.add_run(exp['opis'])
                    self._set_font(opis_run)
                doc.add_paragraph()
        
        # Wykształcenie
        if 'wyksztalcenie' in data and data['wyksztalcenie']:
            self._add_section(doc, 'Wykształcenie', '🎓')
            for edu in data['wyksztalcenie']:
                p = doc.add_paragraph()
                # Kierunek i stopień
                degree = p.add_run(f"{edu['stopien']} - {edu['kierunek']}\n")
                self._set_font(degree, font_size=Pt(12), bold=True)
                # Uczelnia
                uni = p.add_run(f"{edu['uczelnia']}\n")
                self._set_font(uni, font_size=Pt(11))
                # Okres
                period = p.add_run(f"{edu['okres']}")
                self._set_font(period, font_size=Pt(10), italic=True)
                doc.add_paragraph()
        
        # Umiejętności
        if 'umiejetnosci' in data and data['umiejetnosci']:
            self._add_section(doc, 'Umiejętności', '⚙️')
            skills_text = ' • '.join(data['umiejetnosci'])
            p_skills = doc.add_paragraph()
            skills_run = p_skills.add_run(skills_text)
            self._set_font(skills_run)
            doc.add_paragraph()
        
        # Języki
        if 'jezyki' in data and data['jezyki']:
            self._add_section(doc, 'Języki', '🌍')
            for lang in data['jezyki']:
                lang_text = f"{lang['jezyk']}: {lang['poziom']}"
                p_lang = doc.add_paragraph(lang_text, style='List Bullet')
                for run in p_lang.runs:
                    self._set_font(run)
            doc.add_paragraph()
        
        # Zainteresowania
        if 'zainteresowania' in data and data['zainteresowania']:
            self._add_section(doc, 'Zainteresowania', '🎯')
            interests_text = ' • '.join(data['zainteresowania'])
            p_interests = doc.add_paragraph()
            interests_run = p_interests.add_run(interests_text)
            self._set_font(interests_run)
        
        # Stopka
        doc.add_paragraph()
        footer = doc.add_paragraph()
        footer_text = footer.add_run(
            f"Wyrażam zgodę na przetwarzanie moich danych osobowych zgodnie z RODO.\n"
            f"CV wygenerowane przez cv-konwerter.pl - {datetime.now().strftime('%d.%m.%Y')}"
        )
        self._set_font(footer_text, font_size=Pt(8), italic=True)
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Применяем Calibri ко всему документу
        self._apply_calibri_to_all(doc)
        
        # Сохранение
        filename = f"CV_{data['nazwisko']}_{data['imie']}_Klasyczny.docx"
        output_path = self.output_dir / filename
        doc.save(str(output_path))
        
        return output_path
    
    def generate_nowoczesny(self, data: Dict) -> Path:
        """Сгенерировать современное CV с цветным дизайном"""
        doc = Document()
        
        # Настройка отступов
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.4)
            section.bottom_margin = Inches(0.4)
            section.left_margin = Inches(0.6)
            section.right_margin = Inches(0.6)
        
        # HEADER - Имя и Stanowisko (синий фон)
        header = doc.add_paragraph()
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name_run = header.add_run(f"{data['imie']} {data['nazwisko']}\n")
        self._set_font(name_run, font_size=Pt(28), bold=True, color=RGBColor(255, 255, 255))
        
        if 'stanowisko' in data:
            stanowisko_run = header.add_run(data['stanowisko'])
            self._set_font(stanowisko_run, font_size=Pt(14), color=RGBColor(230, 230, 230))
        
        # Добавляем фон (через shading)
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), '1F4E79')  # Синий фон
        header._p.get_or_add_pPr().append(shading_elm)
        
        doc.add_paragraph()
        
        # Контактная информация (иконки + текст)
        contact = doc.add_paragraph()
        contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Формируем адрес
        address_parts = []
        if 'adres' in data and data['adres']:
            address_parts.append(data['adres'])
        if 'kod_pocztowy' in data and data['kod_pocztowy']:
            address_parts.append(data['kod_pocztowy'])
        if 'miasto' in data and data['miasto']:
            address_parts.append(data['miasto'])
        full_address = ', '.join(address_parts) if address_parts else data.get('miasto', '')
        
        contact_text = f"📧 {data['email']}  •  📱 {data['telefon']}  •  📍 {full_address}"
        contact_run = contact.add_run(contact_text)
        self._set_font(contact_run, font_size=Pt(10), color=RGBColor(80, 80, 80))
        
        doc.add_paragraph()
        
        # O sobie (Цитата с цветной линией)
        if 'o_sobie' in data and data['o_sobie']:
            quote = doc.add_paragraph()
            quote_run = quote.add_run(f'"{data["o_sobie"]}"')
            self._set_font(quote_run, font_size=Pt(11), italic=True, color=RGBColor(60, 60, 60))
            quote.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph()
        
        # Doświadczenie zawodowe
        if 'doswiadczenie' in data and data['doswiadczenie']:
            exp_header = doc.add_paragraph()
            exp_title = exp_header.add_run('💼 DOŚWIADCZENIE ZAWODOWE')
            self._set_font(exp_title, font_size=Pt(14), bold=True, color=RGBColor(31, 78, 121))
            
            for exp in data['doswiadczenie']:
                p = doc.add_paragraph()
                # Stanowisko (pogrubione, niebieskie)
                job_title = p.add_run(f"{exp['stanowisko']}\n")
                self._set_font(job_title, font_size=Pt(12), bold=True, color=RGBColor(31, 78, 121))
                
                # Firma i okres (szare)
                company = p.add_run(f"{exp['firma']}  •  {exp['okres']}\n")
                self._set_font(company, font_size=Pt(10), color=RGBColor(100, 100, 100))
                
                # Opis
                if 'opis' in exp and exp['opis']:
                    desc = p.add_run(exp['opis'])
                    self._set_font(desc, font_size=Pt(10))
                
                doc.add_paragraph()
        
        # Wykształcenie
        if 'wyksztalcenie' in data and data['wyksztalcenie']:
            edu_header = doc.add_paragraph()
            edu_title = edu_header.add_run('🎓 WYKSZTAŁCENIE')
            edu_title.font.size = Pt(14)
            edu_title.font.bold = True
            edu_title.font.color.rgb = RGBColor(31, 78, 121)
            
            for edu in data['wyksztalcenie']:
                p = doc.add_paragraph()
                degree = p.add_run(f"{edu['stopien']} - {edu['kierunek']}\n")
                degree.font.bold = True
                degree.font.size = Pt(12)
                degree.font.color.rgb = RGBColor(31, 78, 121)
                
                uni = p.add_run(f"{edu['uczelnia']}  •  {edu['okres']}\n")
                uni.font.size = Pt(10)
                uni.font.color.rgb = RGBColor(100, 100, 100)
                
                doc.add_paragraph()
        
        # Umiejętności (w kolumnach, z цветными тегами)
        if 'umiejetnosci' in data and data['umiejetnosci']:
            skills_header = doc.add_paragraph()
            skills_title = skills_header.add_run('⚙️ UMIEJĘTNOŚCI')
            self._set_font(skills_title, font_size=Pt(14), bold=True, color=RGBColor(31, 78, 121))
            
            # Создаем таблицу для красивого отображения
            table = doc.add_table(rows=1, cols=3)
            table.style = 'Light Grid Accent 1'
            
            cells = table.rows[0].cells
            for i, skill in enumerate(data['umiejetnosci'][:9]):  # Максимум 9
                cell_idx = i % 3
                if i > 0 and cell_idx == 0:
                    cells = table.add_row().cells
                cells[cell_idx].text = f"• {skill}"
                # Устанавливаем шрифт для ячеек
                for paragraph in cells[cell_idx].paragraphs:
                    for run in paragraph.runs:
                        self._set_font(run)
            
            doc.add_paragraph()
        
        # Języki
        if 'jezyki' in data and data['jezyki']:
            lang_header = doc.add_paragraph()
            lang_title = lang_header.add_run('🌍 JĘZYKI')
            self._set_font(lang_title, font_size=Pt(14), bold=True, color=RGBColor(31, 78, 121))
            
            for lang in data['jezyki']:
                p = doc.add_paragraph()
                lang_name = p.add_run(f"{lang['jezyk']}: ")
                self._set_font(lang_name, bold=True)
                lang_level = p.add_run(lang['poziom'])
                self._set_font(lang_level, color=RGBColor(31, 78, 121))
        
        # Stopka
        doc.add_paragraph()
        footer = doc.add_paragraph()
        footer_text = footer.add_run(
            f"Wyrażam zgodę na przetwarzanie moich danych osobowych zgodnie z RODO.\n"
            f"CV wygenerowane przez cv-konwerter.pl - {datetime.now().strftime('%d.%m.%Y')}"
        )
        self._set_font(footer_text, font_size=Pt(8), italic=True, color=RGBColor(120, 120, 120))
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Применяем Calibri ко всему документу
        self._apply_calibri_to_all(doc)
        
        # Сохранение
        filename = f"CV_{data['nazwisko']}_{data['imie']}_Nowoczesny.docx"
        output_path = self.output_dir / filename
        doc.save(str(output_path))
        
        return output_path
    
    def generate_list_motywacyjny(self, data: Dict, cv_data: Dict) -> Path:
        """Сгенерировать сопроводительное письмо"""
        doc = Document()
        
        # Настройка отступов
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        # Dane nadawcy
        sender = doc.add_paragraph()
        sender_runs = [
            sender.add_run(f"{cv_data['imie']} {cv_data['nazwisko']}\n"),
            sender.add_run(f"{cv_data['miasto']}\n"),
            sender.add_run(f"{cv_data['telefon']}\n"),
            sender.add_run(f"{cv_data['email']}\n")
        ]
        for run in sender_runs:
            self._set_font(run)
        sender.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        doc.add_paragraph()
        
        # Data
        date = doc.add_paragraph()
        date_run = date.add_run(f"{data.get('miasto', cv_data['miasto'])}, {datetime.now().strftime('%d.%m.%Y')}")
        self._set_font(date_run)
        date.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        doc.add_paragraph()
        doc.add_paragraph()
        
        # Tytuł
        title = doc.add_paragraph()
        title_run = title.add_run("List motywacyjny")
        self._set_font(title_run, font_size=Pt(16), bold=True)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        
        # Treść
        if 'tresc' in data and data['tresc']:
            doc.add_paragraph(data['tresc'])
        else:
            # Domyślny szablon
            doc.add_paragraph("Szanowni Państwo,")
            doc.add_paragraph()
            
            p1 = doc.add_paragraph(
                f"Z zainteresowaniem przeczytałem/am ogłoszenie o pracę na stanowisku "
                f"{cv_data.get('stanowisko', '[Stanowisko]')}. "
                f"Moje doświadczenie zawodowe oraz umiejętności doskonale pasują do "
                f"wymagań opisanych w ogłoszeniu."
            )
            
            doc.add_paragraph()
            
            p2 = doc.add_paragraph(
                f"Posiadam doświadczenie w pracy zawodowej, które pozwoliło mi rozwinąć "
                f"umiejętności niezbędne na tym stanowisku. Jestem osobą zaangażowaną, "
                f"komunikatywną i gotową do podjęcia nowych wyzwań."
            )
            
            doc.add_paragraph()
            
            p3 = doc.add_paragraph(
                f"Będę wdzięczny/a za rozpatrzenie mojej kandydatury. "
                f"Jestem dostępny/a w dowolnym terminie na rozmowę kwalifikacyjną."
            )
            
            doc.add_paragraph()
            doc.add_paragraph("Z poważaniem,")
            
            signature = doc.add_paragraph()
            sig_run = signature.add_run(f"{cv_data['imie']} {cv_data['nazwisko']}")
            sig_run.font.bold = True
        
        # Применяем Calibri ко всему документу
        self._apply_calibri_to_all(doc)
        
        # Сохранение
        filename = f"List_Motywacyjny_{cv_data['nazwisko']}_{cv_data['imie']}.docx"
        output_path = self.output_dir / filename
        doc.save(str(output_path))
        
        return output_path


# Примеры данных для тестирования
SAMPLE_CV_DATA = {
    'imie': 'Jan',
    'nazwisko': 'Kowalski',
    'email': 'jan.kowalski@example.com',
    'telefon': '+48 123 456 789',
    'miasto': 'Warszawa',
    'stanowisko': 'Junior Python Developer',
    'o_sobie': 'Jestem młodym, ambitnym programistą z pasją do tworzenia aplikacji webowych. Posiadam solidne podstawy programowania w Pythonie oraz doświadczenie w pracy z frameworkiem Flask.',
    'doswiadczenie': [
        {
            'stanowisko': 'Stażysta - Python Developer',
            'firma': 'TechCorp Sp. z o.o.',
            'okres': '2023 - 2024',
            'opis': 'Rozwój aplikacji webowych w Flask, praca z bazami danych SQL, code review, praca w zespole Agile.'
        }
    ],
    'wyksztalcenie': [
        {
            'uczelnia': 'Politechnika Warszawska',
            'kierunek': 'Informatyka',
            'stopien': 'Inżynier',
            'okres': '2019 - 2023'
        }
    ],
    'umiejetnosci': ['Python', 'Flask', 'SQL', 'Git', 'Docker', 'HTML/CSS', 'JavaScript'],
    'jezyki': [
        {'jezyk': 'Polski', 'poziom': 'Ojczysty'},
        {'jezyk': 'Angielski', 'poziom': 'B2'},
        {'jezyk': 'Niemiecki', 'poziom': 'A2'}
    ],
    'zainteresowania': ['Programowanie', 'Sztuczna inteligencja', 'Fotografia', 'Podróże']
}


if __name__ == '__main__':
    # Test generowania
    generator = CVGenerator()
    
    print("🔄 Generowanie testowego CV...")
    cv_path = generator.generate_klasyczny(SAMPLE_CV_DATA)
    print(f"✅ CV wygenerowane: {cv_path}")
    
    print("\n🔄 Generowanie testowego listu motywacyjnego...")
    letter_path = generator.generate_list_motywacyjny({}, SAMPLE_CV_DATA)
    print(f"✅ List motywacyjny wygenerowany: {letter_path}")
